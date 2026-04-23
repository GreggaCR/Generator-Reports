from tkinter import *
from tkinter import ttk
from tkinter import filedialog, messagebox
import os
from datetime import datetime
import re
from docx import Document
from docx.shared import Pt
import sys


#функция для сборки прилодения вместе с иконкой
def resource_path(relative_path):
    """ Получает абсолютный путь к ресурсам, работает для dev и для PyInstaller """
    try:
        # PyInstaller создает временную папку _MEIxxxx
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


# --- СЛОВАРЬ ПАДЕЖЕЙ ДЛЯ ШАБЛОНА ---
PRACTICE_CASES = {
    "Учебная практика — Научно-исследовательская работа(получение первичных навыков научно-исследовательской работы)": "научно-исследовательской работе (получению первичных навыков научно-исследовательской работы)",
    "Производственная практика — Технологическая(проектно-технологическая) практика": "технологической (проектно-технологической) практике",
    "Производственная практика — Педагогическая практика": "педагогической практике",
    "Производственная практика — Научно-исследовательская работа": "научно-исследовательской работе",
    "Производственная практика — Преддипломная практика": "преддипломной практике"
}

def finish():
    root.destroy()

def on_date_key(event):
    entry = event.widget
    if event.keysym in ("BackSpace", "Tab", "Left", "Right"): return
    if not event.char.isdigit(): return "break"
    text = entry.get()
    if len(text) >= 10: return "break"
    if len(text) == 2 or len(text) == 5: entry.insert(END, ".")

def setup_placeholder(entry, placeholder):
    entry.insert(0, placeholder)
    entry.config(foreground="grey")
    def on_focus_in(event):
        if entry.get() == placeholder:
            entry.delete(0, END)
            entry.config(foreground="black")
    def on_focus_out(event):
        if entry.get() == "":
            entry.insert(0, placeholder)
            entry.config(foreground="grey")
    entry.bind("<FocusIn>", on_focus_in)
    entry.bind("<FocusOut>", on_focus_out)

def create_file_selector(parent, label_text, placeholder, is_folder=False):
    container = ttk.Frame(parent, padding=[5, 2])
    container.pack(anchor=NW, padx=10)
    lbl = ttk.Label(container, text=label_text, font=("Times New Roman", 11, "bold"), foreground="blue")
    lbl.pack(anchor=NW, padx=10)
    border_frame = Frame(container, highlightbackground="grey", highlightthickness=1, bg="white")
    border_frame.pack(fill='x', padx=10, pady=5)
    entry = ttk.Entry(border_frame, width=50)
    entry.pack(side=LEFT, padx=(5, 5), pady=5)
    setup_placeholder(entry, placeholder)
    btn_select = ttk.Button(border_frame, text="Обзор...", width=10, 
                            command=lambda: select_item(entry, is_folder, placeholder))
    btn_select.pack(side=LEFT, padx=2)
    def open_path():
        path = entry.get()
        if os.path.exists(path) and path != placeholder:
            os.startfile(path)
        else:
            messagebox.showwarning("Внимание", "Путь не выбран или не существует")
    btn_open = ttk.Button(border_frame, text="👁", width=3, command=open_path)
    btn_open.pack(side=LEFT, padx=2)
    err_lbl = ttk.Label(container, text="", font=("Arial", 8), foreground="red")
    err_lbl.pack(anchor=NW, padx=15)
    return entry, err_lbl

def select_item(entry, is_folder, placeholder):
    path = filedialog.askdirectory() if is_folder else filedialog.askopenfilename(filetypes=(("Word files", "*.docx *.doc"), ("all files", "*.*")))
    if path:
        entry.delete(0, END)
        entry.insert(0, os.path.normpath(path))
        entry.config(foreground="black")

def create_date_selector(parent):
    container = ttk.Frame(parent, padding=[5, 2])
    container.pack(fill='x', padx=10)
    ttk.Label(container, text="Период практики (дд.мм.гггг):", font=("Times New Roman", 11, "bold"), foreground="blue").pack(anchor=NW, padx=10)
    df = ttk.Frame(container); df.pack(anchor=NW, padx=10, pady=5)
    s_date = ttk.Entry(df, width=15); s_date.pack(side=LEFT)
    setup_placeholder(s_date, "дд.мм.гггг"); s_date.bind("<KeyPress>", on_date_key)
    ttk.Label(df, text=" — ").pack(side=LEFT)
    e_date = ttk.Entry(df, width=15); e_date.pack(side=LEFT)
    setup_placeholder(e_date, "дд.мм.гггг"); e_date.bind("<KeyPress>", on_date_key)
    error_label = ttk.Label(container, text="", font=("Arial", 8), foreground="red")
    error_label.pack(anchor=NW, padx=10)
    return s_date, e_date, error_label

def create_type_selector(parent):
    container = ttk.Frame(parent, padding=[5, 2])
    container.pack(fill='x', padx=10)
    lbl = ttk.Label(container, text="Форма отчета для студентов:", font=("Times New Roman", 11, "bold"), foreground="blue")
    lbl.pack(anchor=NW, padx=10)
    vals = ["Бакалавриат", "Магистратура"]
    combobox = ttk.Combobox(container, values=vals, state="readonly", width=30)
    combobox.set(vals[0])
    combobox.pack(anchor=NW, padx=10, pady=5)
    vals_mag = ["Учебная практика — Научно-исследовательская работа(получение первичных навыков научно-исследовательской работы)",
                "Производственная практика — Технологическая(проектно-технологическая) практика",
                "Производственная практика — Педагогическая практика",
                "Производственная практика — Научно-исследовательская работа",
                "Производственная практика — Преддипломная практика"]
    vals_bak =  ["Производственная практика — Технологическая(проектно-технологическая) практика", "Производственная практика — Преддипломная практика"]

    combobox_mag_bak = ttk.Combobox(container, state="readonly",foreground="grey", width=100, font=("Times New Roman", 10, "italic"))
    combobox_mag_bak.set("---Выберите форму отчета---")
    combobox_mag_bak.pack(anchor=NW, padx=10, pady=5)
    err_lbl = ttk.Label(container, text="", font=("Arial", 8), foreground="red")
    err_lbl.pack(anchor=NW, padx=10)
    def on_selected(event):
        err_lbl.config(text="") 
        selected = combobox.get()
        combobox_mag_bak.set("---Выберите форму отчета---")
        if selected == "Магистратура": combobox_mag_bak.configure(values=vals_mag)
        else: combobox_mag_bak.configure(values=vals_bak)
    combobox.bind("<<ComboboxSelected>>", on_selected)
    return combobox, combobox_mag_bak, err_lbl

# --- НОВАЯ ФУНКЦИЯ КОПИРОВАНИЯ СТРОК (ЧТОБЫ НЕ ЕХАЛ ФОРМАТ) ---
def copy_row_format(table, source_row):
    new_row = table.add_row()
    for i, cell in enumerate(source_row.cells):
        # Копируем выравнивание параграфа из ячейки-шаблона
        if cell.paragraphs:
            new_row.cells[i].paragraphs[0].paragraph_format.alignment = cell.paragraphs[0].paragraph_format.alignment
    return new_row

# --- ЛОГИКА ГЕНЕРАЦИИ (ВНЕДРЕНИЕ) ---

def calculate_duration(s, e):
    try:
        d1 = datetime.strptime(s, "%d.%m.%Y")
        d2 = datetime.strptime(e, "%d.%m.%Y")
        delta = d2 - d1
        return f"{delta.days} дн."
    except: return "---"

def check_student_file(folder, name):
    if not os.path.exists(folder): return False
    
    # Очищаем имя от лишних пробелов и берем только фамилию
    clean_name = name.strip()
    surname = clean_name.split()[0].lower()
    
    # Убираем гласные с конца фамилии (останется корень, так проще искать склонения)
    # Например: Двуреченский -> двуреченск, Двуреченского -> двуреченск
    surname_root = re.sub(r'[аеиоуыэюя]$', '', surname) 

    # 1. Сначала ищем по названию файла
    files = [f for f in os.listdir(folder) if f.endswith(".docx") and not f.startswith("~$")]
    for f in files:
        if surname_root in f.lower():
            return True
            
    # 2. Если не нашли, проверяем содержимое (первые 100 абзацев)
    for f in files:
        try:
            path = os.path.join(folder, f)
            doc = Document(path)
            # Берем текст первых 20 абзацев
            header_text = " ".join([p.text for p in doc.paragraphs[:100]]).lower()
            
            if surname_root in header_text:
                return True
        except:
            continue
            
    return False
# --- ПАРСИНГ (ТВОЙ ОРИГИНАЛЬНЫЙ) ---

def parse_vedomost(path):
    data = {"group": "Не найдена", "course": "Не найден", "students": [], "semester": "Не найден"}
    try:
        doc = Document(path)
        lines = [re.sub(r'\s+', ' ', p.text).strip() for p in doc.paragraphs if p.text.strip()]
        for i, line in enumerate(lines):
            if "Курс" in line and i > 0 and lines[i - 1].isdigit(): data["course"] = lines[i - 1]
            if "Семестр" in line and i > 0 and lines[i - 1].isdigit(): data["semester"] = lines[i - 1]

        fio_pattern = re.compile(r'^[А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ][а-яё]+){1,2}$')
        book_pattern = re.compile(r'\d+\/[А-ЯA-Za-z0-9]+')
        
        fio_list = [f for f in lines if fio_pattern.match(f) and f not in ["Фамилия, имя отчество", "Экзаменатор", "Дисциплина"]]
        book_list = []
        for l in lines: book_list.extend(book_pattern.findall(l))
        
        for i in range(min(len(fio_list), len(book_list))):
            data["students"].append({"name": fio_list[i], "grade": "зач"})
    except Exception as e: print(e)
    return data

def show_preview(parsed_data):
    preview_win = Toplevel(root)
    preview_win.title("Предпросмотр данных")
    preview_win.geometry("800x500")
    
    info_frame = ttk.Frame(preview_win, padding=10)
    info_frame.pack(fill=X)
    ttk.Label(info_frame, text=f"Семестр: {parsed_data['semester']} | Курс: {parsed_data['course']}", font=("Arial", 10, "bold")).pack(side=LEFT)

    tree_frame = ttk.Frame(preview_win, padding=10)
    tree_frame.pack(fill=BOTH, expand=True)
    
    tree = ttk.Treeview(tree_frame, columns=("fio", "status"), show='headings')
    tree.heading("fio", text="ФИО Студента"); tree.heading("status", text="Отчет в папке")
    tree.column("fio", width=400); tree.column("status", width=200, anchor=CENTER)
    
    seen_names = set()
    for s in parsed_data["students"]:
        if s["name"] not in seen_names:
# Очищаем имя от случайных пробелов перед проверкой
            student_full_name = s["name"].strip()
            status = "Найден" if check_student_file(entry_folder.get(), student_full_name) else "ОТСУТСТВУЕТ"
            tree.insert("", END, values=(student_full_name, status))
    
    scrollbar = ttk.Scrollbar(tree_frame, orient=VERTICAL, command=tree.yview)
    tree.configure(yscroll=scrollbar.set)
    tree.pack(side=LEFT, fill=BOTH, expand=True); scrollbar.pack(side=RIGHT, fill=Y)

    btn_frame = ttk.Frame(preview_win, padding=10)
    btn_frame.pack(fill=X)
    # ПРИВЯЗЫВАЕМ ГЕНЕРАЦИЮ К КНОПКЕ
    ttk.Button(btn_frame, text="ГЕНЕРИРОВАТЬ DOCX", command=lambda: [preview_win.destroy(), generate_doc(parsed_data)]).pack(side=RIGHT, padx=5)
def extract_student_experience(folder_path, student_name):
    """
    Функция имитирует работу нейросети: ищет файл студента, 
    выделяет 'Вывод' и формулирует виды работ.
    """
    summary = {
        "place": "Не указано",
        "boss_org": "-",
        "position": "студент",
        "work_types": "информация не найдена в отчете"
    }
    
    if not folder_path or not os.path.exists(folder_path):
        return summary

    # Ищем файл (по фамилии)
    surname = student_name.split()[0].lower()
    target_file = None
    for f in os.listdir(folder_path):
        if surname in f.lower() and f.endswith('.docx'):
            target_file = os.path.join(folder_path, f)
            break
    
    if not target_file:
        return summary

    try:
        s_doc = Document(target_file)
        full_text = "\n".join([p.text for p in s_doc.paragraphs])
        
        # 1. Пытаемся вытащить место практики из текста студента (если не нашли в приказе)
        place_match = re.search(r"Место прохождения(?:\s+практики)?[:\-]\s*(.*)", full_text, re.I)
        if place_match:
            summary["place"] = place_match.group(1).strip()

        # 2. ИНТЕЛЛЕКТУАЛЬНЫЙ ПОИСК РАБОТ (Логика "нейросети")
        # Ищем ключевые слова: Вывод, Заключение, В ходе практики...
        markers = ["В ходе практики", "Вывод", "Заключение", "При этом выполнил", "Мною были выполнены"]
        found_text = ""
        
        # Сначала ищем абзац, начинающийся с маркеров
        for p in s_doc.paragraphs:
            text = p.text.strip()
            if any(m.lower() in text.lower() for m in markers):
                found_text = text
                break
        
        # Если нашли блок текста, чистим его от мусора
        if found_text:
            # Убираем вводные слова "В заключение хочу сказать", "Таким образом" и т.д.
            clean_text = re.sub(r"^(.*?)(?:выполнил|сделал|были выполнены)\s*(?:следующие|такие)?\s*(?:виды)?\s*(?:работ)?[:\-]\s*", "", found_text, flags=re.I)
            # Ограничиваем длину и убираем лишние точки
            summary["work_types"] = clean_text.strip()
            
        # 3. Определение должности (стажер/студент/программист)
        if "стажер" in full_text.lower(): summary["position"] = "стажер"
        elif "программист" in full_text.lower(): summary["position"] = "программист-стажер"

    except Exception as e:
        print(f"Ошибка при анализе файла {student_name}: {e}")
        
    return summary
# --- 1. ГЛУБОКИЙ ПОИСК В ПРИКАЗЕ (по таблице) ---
def find_in_prikaz(prikaz_path, student_name):
    res = {"place": "Кафедра ВМИ", "boss": "-"}
    if not prikaz_path or not os.path.exists(prikaz_path):
        return res

    try:
        doc = Document(prikaz_path)
        surname = student_name.split()[0].lower()
        
        for table in doc.tables:
            for row in table.rows:
                # Ищем фамилию во второй ячейке (обычно там ФИО)
                cell_text = row.cells[1].text.lower()
                if surname in cell_text:
                    res["place"] = row.cells[2].text.strip()
                    # Чистим ФИО руководителя от должностей, если нужно
                    res["boss"] = row.cells[3].text.strip().split(',')[0] 
                    return res
    except:
        pass
    return res

def extract_all_text(doc):
    parts = []

    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text:
                    parts.append(c.text)

    return [p.strip() for p in parts if p.strip()]


def find_conclusion_block(lines):
    """
    Берём ПОСЛЕДНЕЕ 'Заключение', но начинаем ПОСЛЕ шапки
    """
    indices = [
        i for i, t in enumerate(lines)
        if "заключение" in t.lower()
        and "руководител" not in t.lower()
    ]

    if not indices:
        return 0

    return indices[-1]


import re

def compress_to_2_sentences(text):
    """
    ГОСТ-стиль: сухой академический вывод без воды
    """

    text = re.sub(r"\s+", " ", text).strip()

    # убираем заголовок
    text = re.sub(r"^(заключение\s*)", "", text, flags=re.I)

    sentences = re.split(r'[.!?]', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]

    # только "официальные" глаголы результата
    verbs = [
        "разработал", "разработана", "разработано",
        "реализовал", "реализована", "реализовано",
        "создал", "создана", "создано",
        "выполнил", "выполнены", "выполнена",
        "провел", "проведена", "проведено",
        "исследовал", "проанализировал",
        "освоил", "изучил", "настроил", "протестировал"
    ]

    def to_gost(sentence):
        sentence = re.sub(r"\b(я|мною|мной|мы|нами)\b", "", sentence, flags=re.I)
        sentence = re.sub(r"\s+", " ", sentence).strip()

        words = sentence.split()

        # ищем первый глагол
        start = None
        for i, w in enumerate(words):
            w_low = w.lower().strip(",.;:-")

            if any(v in w_low for v in verbs):
                start = i
                break

        if start is not None:
            core = " ".join(words[start:])

            # убираем хвосты (иголки)
            core = re.split(r",|;|—|-|\(", core)[0]

            return core.strip()

        return sentence

    result = []

    for s in sentences:
        cleaned = to_gost(s)

        # фильтр мусора
        if len(cleaned) < 15:
            continue

        if any(x in cleaned.lower() for x in [
            "опыт", "расширило понимание", "личные качества"
        ]):
            continue

        result.append(cleaned)

        if len(result) == 2:
            break

    if not result:
        return "выполнены задачи практики и достигнуты цели исследования"

    return ". ".join(result) + "."

def analyze_student_report(student_name, folder_path):
    result = {
        "work_types": "выполнил программу практики",
        "position": "стажер",
        "boss_org": "Не указан"
    }

    if not folder_path or not os.path.exists(folder_path):
        return result

    surname = student_name.split()[0].lower().strip()

    try:
        for file_name in os.listdir(folder_path):
            if not (file_name.endswith(".docx") and not file_name.startswith("~$")):
                continue

            path = os.path.join(folder_path, file_name)
            doc = Document(path)

            header = " ".join(p.text for p in doc.paragraphs[:30]).lower()
            if surname not in header:
                continue

            full_text = "\n".join(p.text for p in doc.paragraphs)

            # --- руководитель ---
            boss_match = re.search(
                r"руководител[ья].{0,40}[:\-]\s*([А-ЯЁ][а-яё\s.]+)",
                full_text,
                re.I
            )
            if boss_match:
                result["boss_org"] = boss_match.group(1).strip()

            # --- полный текст ---
            lines = extract_all_text(doc)

            start = find_conclusion_block(lines)

            conclusion_text = lines[start:]

            # стоп на литературе
            filtered = []
            for line in conclusion_text:
                low = line.lower()

                if "список литературы" in low:
                    break
                if "заключение руководителя" in low:
                    break

                if any(x in low for x in [
                    "фгбоу",
                    "кубгу",
                    "направление подготовки",
                    "профиль",
                    "студента"
                ]):
                    continue

                filtered.append(line)

            raw_text = " ".join(filtered)

            # 🔥 СЖАТИЕ В 1–2 ПРЕДЛОЖЕНИЯ
            result["work_types"] = compress_to_2_sentences(raw_text)

            return result

    except Exception as e:
        print(f"Ошибка анализа {student_name}: {e}")

    return result
# --- 3. УНИВЕРСАЛЬНАЯ ГЕНЕРАЦИЯ (Все шаблоны) ---
def generate_doc(parsed_data):
    try:
        tpl_path = entry_template.get()
        doc = Document(tpl_path)
        
        selected_type = combo_sub_type.get()
        is_nir = "Научно-исследовательская" in selected_type
        
        # --- ЛОГИКА УЧЕБНОГО ГОДА ---
        from datetime import datetime
        start_date_str = entry_start_date.get()
        try:
            start_dt = datetime.strptime(start_date_str, "%d.%m.%Y")
            p_year = start_dt.year
            academic_year = f"{p_year}-{p_year+1}" if start_dt.month >= 9 else f"{p_year-1}-{p_year}"
        except:
            academic_year = "---"

        # --- ОПРЕДЕЛЕНИЕ ТИПА ШАБЛОНА ---
        all_text = " ".join([p.text for p in doc.paragraphs])
        # Для преддиплома ищем наличие обеих меток в тексте параграфов
        is_preddiplom = any("[Оценка]" in p.text and "[ФИО студента]" in p.text for p in doc.paragraphs)
        is_bachelor_std = "[ВИДА_РАБОТ]" in all_text and not is_preddiplom
        is_magistrate_std = "[ФИО СТУДЕНТА] — [Оценка]" in all_text

        # --- ПОДГОТОВКА ДАННЫХ ---
        duration = calculate_duration(start_date_str, entry_end_date.get())
        practice_name = selected_type.capitalize() if is_preddiplom else PRACTICE_CASES.get(selected_type, selected_type)
        
        subs = {
            "[ТИП ПРАКТИКИ]": practice_name,
            "[УЧЕБНЫЙ ГОД]": academic_year,
            "[СРОК ПРОХОЖДЕНИЯ]": duration,
            "[НАПРАВЛЕНИЕ]": "02.03.01 Математика и компьютерные науки" if "бакалавр" in tpl_path.lower() else "02.04.01 Математика и компьютерные науки",
            "[КУРС]": parsed_data.get('course', '---'),
            "[НОМЕР ГРУППЫ]": entry_grp.get(),
            "[ДАТА_НАЧ]": start_date_str,
            "[ДАТА_КОН]": entry_end_date.get(),
        }
        
        for p in doc.paragraphs:
            for k, v in subs.items():
                if k in p.text:
                    p.text = p.text.replace(k, str(v))
            
            if "направления" in p.text.lower() or "направлению" in p.text.lower():
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
            # 2. ДОБАВЛЕНО: Проверка для учебного года
            if "уч. год" in p.text.lower() or academic_year in p.text:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

        # --- РАБОТА С ТАБЛИЦАМИ ---
        if not doc.tables:
            raise Exception("В шаблоне не найдены таблицы")
            
        t0 = doc.tables[0]
        # Используем .strip() для очистки текста ячейки от лишних пробелов при поиске метки
        row_tpl0 = next((r for r in t0.rows if any("[ФИО студента]" in c.text for c in r.cells)), None)
        
        if row_tpl0 is None:
            raise Exception("В первой таблице не найдена строка с меткой [ФИО студента]")

        t1 = doc.tables[1] if len(doc.tables) > 1 else None
        row_tpl1 = None
        if t1:
            row_tpl1 = next((r for r in t1.rows if len(r.cells) > 1 and any("[ФИО студента]" in c.text for c in r.cells)), None)

        summary_list = []
        counts = {"kaf": 0, "org": 0}

        for s in parsed_data['students']:
            info_pr = find_in_prikaz(entry_pr.get(), s['name'])
            info_rep = analyze_student_report(s['name'], entry_folder.get())
            is_kaf = any(x in info_pr['place'] for x in ["Кафедра", "КубГУ", "ВМИ"]) or info_pr['place'] == "Не найдено" or is_nir

            if is_bachelor_std or is_preddiplom:
                counts["kaf"] += 1
                nr = copy_row_format(t0, row_tpl0)
                nr.cells[0].text, nr.cells[1].text, nr.cells[2].text = str(counts["kaf"]), s['name'], info_pr['place']
                
                if is_preddiplom:
                    # В преддипломном таблица может иметь меньше столбцов, проверяем индекс
                    if len(nr.cells) > 3:
                        nr.cells[3].text = info_pr['boss']
                else:
                    if len(nr.cells) > 4:
                        nr.cells[3].text = "-" if is_kaf else (info_rep['boss_org'] if info_rep['boss_org'] != "Не указан" else info_pr['boss'])
                        nr.cells[4].text = "Лахтина А.А."
                
                has_f = check_student_file(entry_folder.get(), s['name'])
                grade = "отлично" if has_f else "неявка"
                
                if is_preddiplom:
                    summary_list.append({"name": s['name'], "text": f" в процессе прохождения {practice_name} практики {info_rep['work_types']} («{grade}»)"})
                else:
                    summary_list.append(f"{s['name']} – {info_pr['place']}, {info_rep['position']}. При этом выполнил следующие виды работ: {info_rep['work_types']}")
            
            else:
                # Обычные магистры (две таблицы)
                if is_kaf:
                    counts["kaf"] += 1
                    nr = copy_row_format(t0, row_tpl0)
                    # ИСПРАВЛЕНО: Точечная замена метки [ПУНКТ], чтобы не стирать содержимое
                    for cell in nr.cells:
                        nr.cells[0].text = str(counts["kaf"])
                        nr.cells[1].text = s['name']
                        if len(nr.cells) > 2: nr.cells[2].text = info_pr['boss']
                        if "[ПУНКТ]" in cell.text:
                            cell.text = cell.text.replace("[ПУНКТ]", str(counts["kaf"]))
                    
                    # Заполняем остальные данные
                    if len(nr.cells) > 2:
                        nr.cells[1].text = s['name']
                        nr.cells[2].text = info_pr['boss']
                
                elif t1 and row_tpl1:
                    counts["org"] += 1
                    nr = copy_row_format(t1, row_tpl1)
                    # Аналогично для второй таблицы, если там есть метка [ПУНКТ]
                    for cell in nr.cells:
                        if "[ПУНКТ]" in cell.text:
                            cell.text = cell.text.replace("[ПУНКТ]", str(counts["org"]))
                        elif "[№]" in cell.text: # На случай, если метка называется иначе
                            cell.text = cell.text.replace("[№]", str(counts["org"]))

                    nr.cells[0].text = str(counts["org"]) # Если там просто номер без метки
                    nr.cells[1].text = s['name']
                    if len(nr.cells) > 4:
                        nr.cells[2].text = info_pr['place']
                        nr.cells[3].text = info_pr['boss']
                        nr.cells[4].text = "Фоменко С.И."

        # --- ОЧИСТКА ---
        if row_tpl0: t0._element.remove(row_tpl0._element)
        if t1 and row_tpl1: t1._element.remove(row_tpl1._element)
        
        if is_magistrate_std and (is_nir or counts["org"] == 0) and t1:
            try:
                parent = t1._element.getparent()
                tbl_idx = list(parent).index(t1._element)
                if tbl_idx > 0:
                    parent.remove(parent[tbl_idx - 1])
                parent.remove(t1._element)
            except:
                pass 

        # --- ФИНАЛЬНЫЙ ТЕКСТ ---
        for p in doc.paragraphs:
            # Преддипломный блок
            if is_preddiplom and "[ФИО студента]" in p.text and "[Оценка]" in p.text:
                p.text = ""
                for item in summary_list:
                    r1 = p.add_run("Студент " + item["name"]); r1.font.bold = True
                    r2 = p.add_run(item["text"] + ".\n")
                    for r in [r1, r2]: r.font.name, r.font.size = 'Times New Roman', Pt(14)

            # Бакалаврский блок
            elif is_bachelor_std and "[ВИДА_РАБОТ]" in p.text:
                p.text = ""
                for line in summary_list:
                    name_part, rest_part = line.split(" – ", 1) if " – " in line else (line, "")
                    r1 = p.add_run(name_part); r1.font.bold = True
                    r2 = p.add_run(" – " + rest_part + ".\n")
                    for r in [r1, r2]: r.font.name, r.font.size = 'Times New Roman', Pt(14)

            # Магистерский стандарт
            elif is_magistrate_std and "[ФИО СТУДЕНТА] — [Оценка]" in p.text:
                p.text = ""
                for s in parsed_data['students']:
                    grade = "зач" if check_student_file(entry_folder.get(), s['name']) else "неявка"
                    run = p.add_run(f"{s['name']} — {grade}.\n")
                    run.font.name, run.font.size = 'Times New Roman', Pt(14)

        out_path = os.path.join(os.path.dirname(tpl_path), f"ИТОГ_{entry_grp.get()}.docx")
        doc.save(out_path)
        os.startfile(out_path)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Сбой: {str(e)}")
def start_gen():

    for lbl in [err_ved, err_pr, err_fld, err_tpl, date_error_lbl, err_type]:
        lbl.config(text="")

    errors = False

    # 1. Проверка заполнения всех обязательных полей
    if not entry_ved.get() or entry_ved.get() == "Укажите файл ведомости...":
        err_ved.config(text="⚠️ Укажите файл ведомости")
        errors = True
    if not entry_pr.get() or entry_pr.get() == "Укажите файл приказа...":
        err_pr.config(text="⚠️ Укажите файл приказа")
        errors = True
    if not entry_folder.get() or entry_folder.get() == "Выберите папку...":
        err_fld.config(text="⚠️ Укажите папку с отчётами студентов")
        errors = True
    if not entry_template.get() or entry_template.get() == "Укажите файл шаблона (.docx)...":
        err_tpl.config(text="⚠️ Укажите файл шаблона")
        errors = True
    if not entry_grp.get().strip():
        # Добавляем временную красную метку под группой (можно создать отдельный err_grp, но чтобы минимально менять — используем messagebox)
        messagebox.showwarning("Внимание", "Не указан номер группы!")
        errors = True

    # 2. Проверка существования файлов/папки
    if entry_ved.get() and not os.path.exists(entry_ved.get()):
        err_ved.config(text="⚠️ Файл ведомости не найден")
        messagebox.showerror("Ошибка", "Файл ведомости не существует!")
        errors = True
    if entry_pr.get() and not os.path.exists(entry_pr.get()):
        err_pr.config(text="⚠️ Файл приказа не найден")
        messagebox.showerror("Ошибка", "Файл приказа не существует!")
        errors = True
    if entry_folder.get() and not os.path.exists(entry_folder.get()):
        err_fld.config(text="⚠️ Папка с отчётами не найдена")
        if not entry_grp.get().strip():
            err_grp.config(text="⚠️ Укажите группу")
            errors = True
    if entry_template.get() and not os.path.exists(entry_template.get()):
        err_tpl.config(text="⚠️ Файл шаблона не найден")
        messagebox.showerror("Ошибка", "Файл шаблона итогового отчета не существует!")
        errors = True

    # 3. Проверка дат
    s_date = entry_start_date.get().strip()
    e_date = entry_end_date.get().strip()

    if s_date == "дд.мм.гггг" or e_date == "дд.мм.гггг" or not s_date or not e_date:
        date_error_lbl.config(text="⚠️ Укажите обе даты")
        messagebox.showwarning("Внимание", "Необходимо указать период практики (обе даты)")
        errors = True
    else:
        try:
            d1 = datetime.strptime(s_date, "%d.%m.%Y")
            d2 = datetime.strptime(e_date, "%d.%m.%Y")
            
            if d1 > d2:
                date_error_lbl.config(text="⚠️ Дата начала не может быть позже даты окончания")
                messagebox.showwarning("Ошибка дат", "Дата начала практики не может быть позже даты окончания!")
                errors = True
            elif d2 > datetime.now():
                date_error_lbl.config(text="⚠️ Дата окончания не может быть в будущем")
                messagebox.showwarning("Ошибка дат", "Дата окончания практики не может быть в будущем!")
                errors = True
            # Дополнительно: проверка разумного диапазона (не больше 6 месяцев, например)
            elif (d2 - d1).days > 200:
                date_error_lbl.config(text="⚠️ Слишком большой период практики")
                messagebox.showwarning("Предупреждение", "Период практики больше 200 дней — проверьте даты")
                # errors = False  # только предупреждение, не блокируем
        except ValueError:
            date_error_lbl.config(text="⚠️ Неверный формат даты или несуществующая дата")
            messagebox.showerror("Ошибка", "Дата введена некорректно!\nИспользуйте формат дд.мм.гггг")
            errors = True

    # 4. Проверка типа практики
    if combo_sub_type.get() == "---Выберите форму отчета---":
        err_type.config(text="⚠️ Выберите форму отчета")
        messagebox.showwarning("Внимание", "Не выбрана форма отчета (тип практики)")
        errors = True

    if errors:
        return  # ← прерываем выполнение, если есть ошибки

    # ====================== ОРИГИНАЛЬНАЯ ЛОГИКА (без изменений) ======================
    path_ved = entry_ved.get()

    # Парсим ведомость
    results = parse_vedomost(path_ved)
    
    # Добавляем в данные студента "умные" поля
    folder_reports = entry_folder.get()
    for student in results["students"]:
        analysis = analyze_student_report(student["name"], folder_reports)
        student["work_types"] = analysis["work_types"]
        student["position"] = analysis["position"]

    # Показываем предпросмотр
    show_preview(results)

# --- ИНТЕРФЕЙС (ТВОЙ) ---
root = Tk()
root.title("RepGen")
root.geometry("850x900")
root.iconbitmap(resource_path("web.ico"))

ttk.Label(text="Параметры формирования отчета", font=("Times New Roman", 16)).pack(pady=10)
entry_ved, err_ved = create_file_selector(root, "Ведомость", "Укажите файл ведомости...")
entry_pr, err_pr = create_file_selector(root, "Приказ", "Укажите файл приказа...")
entry_folder, err_fld = create_file_selector(root, "Папка с отчетами студентов", "Выберите папку...", True)
entry_template, err_tpl = create_file_selector(root, "Шаблон итогового отчета", "Укажите файл шаблона (.docx)...")

# Добавил поле Группы, оно нужно
f_grp = ttk.Frame(root); f_grp.pack(anchor=NW, padx=25)
ttk.Label(f_grp, text="Группа:", font=("Times New Roman", 11, "bold"), foreground="blue").pack(side=LEFT)
entry_grp = ttk.Entry(f_grp, width=20); entry_grp.pack(side=LEFT, padx=10)
err_grp = ttk.Label(f_grp, text="", font=("Arial", 8), foreground="red")
err_grp.pack(side=LEFT, padx=10)

entry_start_date, entry_end_date, date_error_lbl = create_date_selector(root)
combo_main_type, combo_sub_type, err_type = create_type_selector(root)

btn_gen = ttk.Button(root, text="ЗАПУСТИТЬ ОБРАБОТКУ", command=start_gen)
btn_gen.pack(pady=20)

root.mainloop()